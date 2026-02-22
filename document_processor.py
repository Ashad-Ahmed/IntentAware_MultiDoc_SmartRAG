import os
import logging
from typing import List, Tuple
from pathlib import Path
from config import settings

logger = logging.getLogger(__name__)


class DocumentProcessor:
    """Extract and chunk text from multiple document formats"""

    @staticmethod
    def extract_text(file_path: str, file_type: str) -> str:
        """
        Extract text from document based on file type

        Args:
            file_path: Path to the document
            file_type: Document type (txt, pdf, docx)

        Returns:
            Extracted text content
        """
        file_type = file_type.lower()

        if file_type == "txt":
            return DocumentProcessor._extract_txt(file_path)
        elif file_type == "pdf":
            return DocumentProcessor._extract_pdf(file_path)
        elif file_type == "docx":
            return DocumentProcessor._extract_docx(file_path)
        else:
            raise ValueError(f"Unsupported file type: {file_type}")

    @staticmethod
    def _extract_txt(file_path: str) -> str:
        """Extract text from .txt files"""
        try:
            with open(file_path, "r", encoding="utf-8") as f:
                return f.read()
        except Exception as e:
            logger.error(f"Error reading TXT file {file_path}: {e}")
            raise

    @staticmethod
    def _extract_pdf(file_path: str) -> str:
        """Extract text from PDF files"""
        try:
            import pdfplumber

            text_content = []
            with pdfplumber.open(file_path) as pdf:
                for page_num, page in enumerate(pdf.pages, 1):
                    text = page.extract_text()
                    if text:
                        # Add page markers for better context
                        text_content.append(f"[PAGE {page_num}]\n{text}")

            return "\n\n".join(text_content)
        except ImportError:
            logger.error("pdfplumber not installed. Install with: pip install pdfplumber")
            raise
        except Exception as e:
            logger.error(f"Error reading PDF file {file_path}: {e}")
            raise

    @staticmethod
    def _extract_docx(file_path: str) -> str:
        """Extract text from DOCX files"""
        try:
            from docx import Document

            doc = Document(file_path)
            text_content = []

            for para in doc.paragraphs:
                if para.text.strip():
                    text_content.append(para.text)

            # Extract text from tables if present
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(cell.text for cell in row.cells)
                    if row_text.strip():
                        text_content.append(row_text)

            return "\n".join(text_content)
        except ImportError:
            logger.error("python-docx not installed. Install with: pip install python-docx")
            raise
        except Exception as e:
            logger.error(f"Error reading DOCX file {file_path}: {e}")
            raise

    @staticmethod
    def chunk_text(
        text: str,
        chunk_words: int = settings.CHUNK_WORDS,
        overlap_words: int = settings.CHUNK_OVERLAP,
    ) -> List[Tuple[int, str]]:
        """
        Split text into overlapping chunks

        Args:
            text: Text to chunk
            chunk_words: Words per chunk
            overlap_words: Overlap between chunks

        Returns:
            List of (start_word_index, chunk_text) tuples
        """
        words = text.split()
        if not words:
            return []

        step = chunk_words - overlap_words
        chunks = []

        for i in range(0, len(words), step):
            chunk_start = i
            chunk_end = min(i + chunk_words, len(words))
            chunk_text = " ".join(words[chunk_start:chunk_end])

            if chunk_text.strip():
                chunks.append((chunk_start, chunk_text))

        return chunks

    @staticmethod
    def validate_file(file_path: str, file_name: str) -> Tuple[bool, str]:
        """
        Validate uploaded file

        Args:
            file_path: Path to uploaded file
            file_name: Original filename

        Returns:
            (is_valid, error_message)
        """
        # Check file exists
        if not os.path.exists(file_path):
            return False, "File not found"

        # Check file size
        file_size = os.path.getsize(file_path)
        if file_size > settings.MAX_FILE_SIZE:
            return (
                False,
                f"File too large. Max size: {settings.MAX_FILE_SIZE / 1024 / 1024:.1f}MB",
            )

        # Check extension
        file_ext = file_name.split(".")[-1].lower()
        if file_ext not in settings.ALLOWED_EXTENSIONS:
            return (
                False,
                f"Unsupported file type. Allowed: {', '.join(settings.ALLOWED_EXTENSIONS)}",
            )

        return True, ""
